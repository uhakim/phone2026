from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_korean_font(run, font_name="맑은 고딕"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text, size=13):
    p = doc.add_paragraph(text)
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(size)
    set_korean_font(r)
    return p


def main():
    base = Path(__file__).resolve().parents[1]
    img_dir = base / "docs" / "guide_assets"
    out_file = base / "docs" / "parent_notice_streamlit_guide.docx"
    out_file.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(11)

    p = doc.add_paragraph("가정통신문")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(20)
    set_korean_font(r)

    p = doc.add_paragraph("출입·스마트기기 관리시스템 사용 안내")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(16)
    set_korean_font(r)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    r1 = p.add_run("학부모님 안녕하십니까?\n")
    r1.bold = True
    set_korean_font(r1)
    r2 = p.add_run(
        "항상 학교 교육활동에 관심과 협조를 보내주셔서 감사드립니다. "
        "우리 학교는 학부모 신청과 학교 승인 절차를 보다 정확하고 편리하게 처리하기 위해 "
        "「출입·스마트기기 관리시스템」을 운영하고 있습니다.\n"
        "아래 안내를 참고하시어 학생의 정문 출입, 휴대전화 및 수업용 태블릿PC 신청을 진행해 주시기 바랍니다."
    )
    set_korean_font(r2)

    add_heading(doc, "1. 시스템 안내")
    for t in [
        "운영 목적: 학생 안전 관리 및 신청·승인 절차의 전산화",
        "신청 항목: ① 휴대전화 소지 허가 ② 수업용 태블릿PC 소지 허가 ③ 정문 출입 허가",
        "처리 방식: 학부모 신청 → 관리자 확인/승인 → 허가서(PDF) 발급",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")
        set_korean_font(p.runs[0])

    add_heading(doc, "2. 접속 방법")
    for t in [
        "인터넷 브라우저(크롬/엣지)에서 학교 안내 주소로 접속합니다.",
        "왼쪽 메뉴에서 「학부모 페이지」를 선택합니다.",
        "학생 인증: 학번(예: 1101), 이름 입력 후 인증하기 버튼을 누릅니다.",
    ]:
        p = doc.add_paragraph(t, style="List Number")
        set_korean_font(p.runs[0])

    img1 = img_dir / "01_main.png"
    if img1.exists():
        p = doc.add_paragraph("화면 1) 시스템 메인 화면")
        set_korean_font(p.runs[0])
        doc.add_picture(str(img1), width=Inches(6.3))

    img2 = img_dir / "05_parent_login_filled.png"
    if img2.exists():
        p = doc.add_paragraph("화면 2) 학부모 페이지 학생 인증 입력 예시(학번 1101, 이름 홍길동)")
        set_korean_font(p.runs[0])
        doc.add_picture(str(img2), width=Inches(6.3))

    add_heading(doc, "3. 신청 절차(학부모)")
    for t in [
        "인증 후 상단 탭에서 신청 항목(휴대전화/태블릿PC/정문출입)을 선택합니다.",
        "필수 항목(사유, 시간 등)을 정확히 입력하고 신청합니다.",
        "신청 현황 탭에서 승인 상태를 확인하고, 필요 시 취소할 수 있습니다.",
        "승인 완료 후 허가서 PDF를 내려받아 활용합니다.",
    ]:
        p = doc.add_paragraph(t, style="List Number")
        set_korean_font(p.runs[0])

    add_heading(doc, "4. 승인 후 허가서 출력 방법(학부모)")
    for t in [
        "학부모 페이지 로그인 후 「신청 현황」 탭으로 이동합니다.",
        "승인 완료(초록 점) 상태의 신청 카드 오른쪽에서 「PDF 출력」 버튼을 누릅니다.",
        "다운로드된 파일이 학교 허가서 양식(PDF)입니다.",
    ]:
        p = doc.add_paragraph(t, style="List Number")
        set_korean_font(p.runs[0])

    img_status = img_dir / "06_parent_status_pdf_button.png"
    if img_status.exists():
        p = doc.add_paragraph("화면 3) 신청 현황 탭에서 PDF 출력 버튼 위치")
        set_korean_font(p.runs[0])
        doc.add_picture(str(img_status), width=Inches(6.3))

    img_pdf = img_dir / "07_phone_permit_output.png"
    if img_pdf.exists():
        p = doc.add_paragraph("화면 4) PDF 출력 결과(휴대전화 허가서 양식)")
        set_korean_font(p.runs[0])
        doc.add_picture(str(img_pdf), width=Inches(6.3))
    add_heading(doc, "5. 관리자 처리 절차(학교용)")
    p = doc.add_paragraph("※ 아래 내용은 학교 내부 처리용 안내입니다.")
    p.runs[0].italic = True
    set_korean_font(p.runs[0])

    img3 = img_dir / "03_admin_approval.png"
    if img3.exists():
        p = doc.add_paragraph("화면 3) 관리자 승인 페이지")
        set_korean_font(p.runs[0])
        doc.add_picture(str(img3), width=Inches(6.3))

    img4 = img_dir / "04_admin_manage.png"
    if img4.exists():
        p = doc.add_paragraph("화면 4) 관리 페이지(학생 명단/설정)")
        set_korean_font(p.runs[0])
        doc.add_picture(str(img4), width=Inches(6.3))

    add_heading(doc, "6. 유의사항")
    for t in [
        "학부모 신청 정보(사유/시간)는 사실에 맞게 작성해 주십시오.",
        "정문 출입 시간 변경이 필요한 경우, 기존 신청을 취소 후 다시 신청해 주십시오.",
        "허가서 관련 문의는 담임교사 또는 생활안전 담당부서로 연락해 주십시오.",
    ]:
        p = doc.add_paragraph(t, style="List Bullet")
        set_korean_font(p.runs[0])

    doc.add_paragraph("")
    p = doc.add_paragraph(f"{date.today().year}년 {date.today().month}월 {date.today().day}일")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_korean_font(p.runs[0])

    p = doc.add_paragraph("동성초등학교장")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)
    set_korean_font(r)

    doc.save(str(out_file))
    print(out_file)


if __name__ == "__main__":
    main()
