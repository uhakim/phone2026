from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


GOOGLE_SHEET_URL = "https://docs.google.com/spreadsheets/d/16QWXBF_HSl0T55JmEJLp46ulJsGuPt7zVp3LAGBZHZM/edit?gid=0#gid=0"
APP_URL = "https://phone2026-jciaey8jvalvktcs3xr6df.streamlit.app/"


def set_korean_font(run, font_name="맑은 고딕"):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_heading(doc, text, size=13):
    p = doc.add_paragraph(text)
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(size)
    set_korean_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    set_korean_font(p.runs[0])


def add_number(doc, text):
    p = doc.add_paragraph(text, style="List Number")
    set_korean_font(p.runs[0])


def add_image_with_caption(doc, image_path: Path, caption: str, width=6.3):
    if not image_path.exists():
        return
    p = doc.add_paragraph(caption)
    set_korean_font(p.runs[0])
    doc.add_picture(str(image_path), width=Inches(width))


def main():
    base = Path(__file__).resolve().parents[1]
    img_dir = base / "docs" / "guide_assets"
    out_dir = base / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_file = out_dir / "담임교사_안내용_시스템_매뉴얼.docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(11)

    p = doc.add_paragraph("담임교사 안내용 매뉴얼")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(20)
    set_korean_font(r)

    p = doc.add_paragraph("출입·스마트기기 관리시스템 운영 절차")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(15)
    set_korean_font(r)

    doc.add_paragraph("")
    p = doc.add_paragraph("1) 시스템 접속 정보")
    p.runs[0].bold = True
    set_korean_font(p.runs[0])
    p = doc.add_paragraph(f"학생/학부모 접속 주소: {APP_URL}")
    set_korean_font(p.runs[0])
    p = doc.add_paragraph(f"정문출입 구글시트 주소: {GOOGLE_SHEET_URL}")
    set_korean_font(p.runs[0])

    add_heading(doc, "1. 전체 운영 흐름")
    add_number(doc, "학부모가 학부모 페이지에서 신청(휴대전화/태블릿PC/정문출입)")
    add_number(doc, "관리자 승인 페이지에서 승인 또는 반려 처리")
    add_number(doc, "승인 완료 시 학부모가 신청 현황 탭에서 PDF 출력")
    add_number(doc, "정문출입 승인 건은 구글시트 명단으로 자동 동기화")
    add_image_with_caption(doc, img_dir / "01_main.png", "화면 1) 시스템 메인 화면")

    add_heading(doc, "2. 학부모 안내 핵심 문구(담임 전달용)")
    add_bullet(doc, "학번과 이름으로 로그인 후 신청서를 작성합니다.")
    add_bullet(doc, "승인 상태는 신청 현황 탭에서 확인할 수 있습니다.")
    add_bullet(doc, "승인 완료 후 PDF 출력 버튼으로 허가서를 내려받습니다.")
    add_bullet(doc, "허가서는 출력하여 학생이 항상 휴대할 수 있도록 지도해 주시기 바랍니다.")
    add_image_with_caption(
        doc,
        img_dir / "05_parent_login_filled.png",
        "화면 2) 학부모 페이지 로그인 입력 예시(학번 1101, 이름 홍길동)",
    )
    add_image_with_caption(doc, img_dir / "02_parent_login.png", "화면 3) 학부모 페이지 로그인 화면(기본)")

    add_heading(doc, "3. 관리자 승인 절차")
    add_number(doc, "좌측 메뉴에서 「관리자 승인 페이지」 선택")
    add_number(doc, "관리자 비밀번호 입력 후 로그인")
    add_number(doc, "대기 신청 목록에서 신청 건 확인")
    add_number(doc, "승인 또는 반려 처리(반려 시 사유 입력)")
    add_number(doc, "처리 결과를 학부모가 신청 현황에서 확인")
    add_image_with_caption(doc, img_dir / "03_admin_approval.png", "화면 4) 관리자 승인 페이지")

    add_heading(doc, "4. 관리 페이지 주요 기능")
    add_bullet(doc, "학생 명단 관리: CSV 일괄 업로드, 개별 추가/삭제, 전체 삭제")
    add_bullet(doc, "승인 모드 설정: 자동 발급 / 승인 필요")
    add_bullet(doc, "학년도 설정: 학년도/시작일 설정(종료일 자동 계산)")
    add_bullet(doc, "문서 관리: 학교장 확인 도장 이미지 업로드")
    add_bullet(doc, "정문 출입 명단: 표 조회 및 구글시트 업로드")
    add_image_with_caption(doc, img_dir / "04_admin_manage.png", "화면 5) 관리 페이지")

    add_heading(doc, "5. 정문출입 구글시트 연동 안내")
    add_bullet(doc, "연동 대상: 승인 완료된 정문출입 신청 건")
    add_bullet(doc, "입력 위치: 시트 A4:M 영역")
    add_bullet(doc, "등교(C~G)는 체크박스 값, 하교(H~L)는 1하교/2하교/3하교 텍스트")
    add_bullet(doc, f"확인 주소: {GOOGLE_SHEET_URL}")

    add_heading(doc, "6. 학부모가 실제로 PDF를 출력하는 위치")
    add_number(doc, "학부모 페이지 로그인")
    add_number(doc, "「신청 현황」 탭 클릭")
    add_number(doc, "승인 완료 카드 오른쪽 「PDF 출력」 버튼 클릭")
    add_image_with_caption(doc, img_dir / "06_parent_status_pdf_button.png", "화면 6) 신청 현황 탭의 PDF 출력 버튼")
    add_image_with_caption(doc, img_dir / "07_phone_permit_output.png", "화면 7) 휴대전화 허가서 출력 양식")
    add_image_with_caption(doc, img_dir / "08_gate_permit_output.png", "화면 8) 정문 출입 허가서 출력 양식")

    add_heading(doc, "7. 장애 대응 체크리스트")
    add_bullet(doc, "로그인 실패: 학번/이름 일치 여부 확인")
    add_bullet(doc, "신청이 보이지 않음: 승인 페이지에서 상태 확인")
    add_bullet(doc, "PDF 양식이 다름: 템플릿 파일 반영 여부 확인")
    add_bullet(doc, "구글시트 미반영: 승인 여부 및 스크립트 URL 설정 확인")
    qr = base / "docs" / "qr_phone2026_app.png"
    add_image_with_caption(doc, qr, "부록) 학부모용 시스템 접속 QR 코드", width=2.0)

    doc.add_paragraph("")
    p = doc.add_paragraph(f"{date.today().year}년 {date.today().month}월 {date.today().day}일")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_korean_font(p.runs[0])

    p = doc.add_paragraph("동성초등학교")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(13)
    set_korean_font(r)

    doc.save(str(out_file))
    print(out_file)


if __name__ == "__main__":
    main()
